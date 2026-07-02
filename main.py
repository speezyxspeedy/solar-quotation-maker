from __future__ import annotations

import json
import sqlite3
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from pydantic import BaseModel, Field
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas


BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = BASE_DIR / "assets" / "logo.png"
DB_PATH = BASE_DIR / "sgse_app.db"

COMPANY_NAME = "SHREE GANESH SOLAR ENERGY"
OFFICE_ADDRESS = "P/N 1764/29 HANUMAN NAGAR BHANDEWADI PARDI NAGPUR 440008"
EMAIL = "rakeshbhude69@gmail.com"
GST_NO = "27DSQPB6025M1ZC"
GST_FOOTER = "27DSQPB6025M"
CONTACT_NO = "+91 7385284731"
BANK_NAME = "INDIAN BANK"
ACCOUNT_HOLDER = "SHREE GANESH SOLAR ENERGY"
CC_ACCOUNT_NO = "8275515270"
IFSC_CODE = "IDIB000R558"
BRANCH = "DESHPANDE LAYOUT, NAGPUR"

PAGE_WIDTH, PAGE_HEIGHT = A4
MARGIN = 36
GREEN = colors.HexColor("#176B3A")
DARK = colors.HexColor("#172033")
MUTED = colors.HexColor("#536170")
GOLD = colors.HexColor("#E0A623")
LINE = colors.HexColor("#C8D3DC")
SOFT_GREEN = colors.HexColor("#EAF4ED")
SOFT_GOLD = colors.HexColor("#FFF4D4")


class Specification(BaseModel):
    item: str = ""
    brand_parameters: str = ""
    detail: str = ""
    specification: str = ""
    quantity: str = ""


class PriceRow(BaseModel):
    sr_no: str = ""
    system_detail: str = ""
    rate: str = ""


class Quotation(BaseModel):
    id: str = ""
    quotation_number: str
    customer_name: str
    location: str
    capacity_kw: str
    quote_date: str
    final_price: str = ""
    panel_brand: str = ""
    inverter_brand: str = ""
    introduction: str = ""
    specifications: List[Specification] = Field(default_factory=list)
    price_header: str = ""
    price_rows: List[PriceRow] = Field(default_factory=list)
    price_note: str = ""
    payment_note: str = ""
    terms_conditions: str = ""
    remark: str = ""


class InvoiceParty(BaseModel):
    name: str = ""
    address: str = ""
    gstin: str = ""
    state: str = ""
    state_code: str = ""


class InvoiceItem(BaseModel):
    sl_no: str = ""
    description: str = ""
    hsn_sac: str = ""
    quantity: float = 0
    unit: str = ""
    rate: float = 0
    amount: float = 0


class InvoiceTaxSummaryRow(BaseModel):
    hsn_sac: str = ""
    taxable_value: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    total_tax: float = 0


class InvoiceBankDetails(BaseModel):
    account_holder: str = ""
    bank_name: str = ""
    account_no: str = ""
    ifsc: str = ""
    branch: str = ""


class Invoice(BaseModel):
    id: str = ""
    invoice_no: str
    invoice_date: str
    seller: InvoiceParty = Field(default_factory=InvoiceParty)
    buyer: InvoiceParty = Field(default_factory=InvoiceParty)
    ship_to: InvoiceParty = Field(default_factory=InvoiceParty)
    delivery_note: str = ""
    payment_terms: str = ""
    dispatch_through: str = ""
    destination: str = ""
    vehicle_no: str = ""
    terms_of_delivery: str = ""
    items: List[InvoiceItem] = Field(default_factory=list)
    use_igst: bool = False
    cgst_rate: float = 9
    sgst_rate: float = 9
    igst_rate: float = 18
    subtotal: float = 0
    cgst: float = 0
    sgst: float = 0
    igst: float = 0
    round_off: float = 0
    grand_total: float = 0
    total_quantity: float = 0
    amount_in_words: str = ""
    tax_summary: List[InvoiceTaxSummaryRow] = Field(default_factory=list)
    bank_details: InvoiceBankDetails = Field(default_factory=InvoiceBankDetails)
    declaration: str = ""


app = FastAPI(title=f"{COMPANY_NAME} - Solar Quotation Generator")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

saved_quotations: list[Quotation] = []


def dump_model(model: BaseModel) -> dict:
    return model.model_dump() if hasattr(model, "model_dump") else model.dict()


def db_connection() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def init_db() -> None:
    with db_connection() as connection:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS invoices (
                id TEXT PRIMARY KEY,
                invoice_no TEXT UNIQUE NOT NULL,
                invoice_date TEXT NOT NULL,
                seller_company_name TEXT,
                seller_address TEXT,
                seller_gstin TEXT,
                seller_state TEXT,
                seller_state_code TEXT,
                buyer_name TEXT,
                buyer_address TEXT,
                buyer_gstin TEXT,
                buyer_state TEXT,
                buyer_state_code TEXT,
                ship_to_name TEXT,
                ship_to_address TEXT,
                ship_to_gstin TEXT,
                ship_to_state TEXT,
                ship_to_state_code TEXT,
                delivery_note TEXT,
                payment_terms TEXT,
                dispatch_through TEXT,
                destination TEXT,
                vehicle_no TEXT,
                terms_of_delivery TEXT,
                subtotal REAL,
                cgst REAL,
                sgst REAL,
                igst REAL,
                round_off REAL,
                grand_total REAL,
                total_quantity REAL,
                amount_in_words TEXT,
                bank_details_json TEXT,
                declaration TEXT,
                raw_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS invoice_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                invoice_id TEXT NOT NULL,
                sl_no TEXT,
                description TEXT,
                hsn_sac TEXT,
                quantity REAL,
                unit TEXT,
                rate REAL,
                amount REAL,
                FOREIGN KEY(invoice_id) REFERENCES invoices(id) ON DELETE CASCADE
            )
            """
        )
        connection.execute("CREATE INDEX IF NOT EXISTS invoice_items_invoice_id_idx ON invoice_items(invoice_id)")
        connection.execute("CREATE INDEX IF NOT EXISTS invoices_invoice_date_idx ON invoices(invoice_date)")


def save_invoice_to_db(invoice: Invoice) -> dict:
    now = datetime.now().isoformat(timespec="seconds")
    if not invoice.id:
        invoice.id = str(uuid4())

    payload = dump_model(invoice)
    with db_connection() as connection:
        existing = connection.execute(
            "SELECT id, created_at FROM invoices WHERE invoice_no = ?",
            (invoice.invoice_no,),
        ).fetchone()
        created_at = existing["created_at"] if existing else now
        if existing and existing["id"] != invoice.id:
            invoice.id = existing["id"]
            payload["id"] = invoice.id

        connection.execute(
            """
            INSERT INTO invoices (
                id, invoice_no, invoice_date,
                seller_company_name, seller_address, seller_gstin, seller_state, seller_state_code,
                buyer_name, buyer_address, buyer_gstin, buyer_state, buyer_state_code,
                ship_to_name, ship_to_address, ship_to_gstin, ship_to_state, ship_to_state_code,
                delivery_note, payment_terms, dispatch_through, destination, vehicle_no, terms_of_delivery,
                subtotal, cgst, sgst, igst, round_off, grand_total, total_quantity, amount_in_words,
                bank_details_json, declaration, raw_json, created_at, updated_at
            ) VALUES (
                :id, :invoice_no, :invoice_date,
                :seller_company_name, :seller_address, :seller_gstin, :seller_state, :seller_state_code,
                :buyer_name, :buyer_address, :buyer_gstin, :buyer_state, :buyer_state_code,
                :ship_to_name, :ship_to_address, :ship_to_gstin, :ship_to_state, :ship_to_state_code,
                :delivery_note, :payment_terms, :dispatch_through, :destination, :vehicle_no, :terms_of_delivery,
                :subtotal, :cgst, :sgst, :igst, :round_off, :grand_total, :total_quantity, :amount_in_words,
                :bank_details_json, :declaration, :raw_json, :created_at, :updated_at
            )
            ON CONFLICT(id) DO UPDATE SET
                invoice_no = excluded.invoice_no,
                invoice_date = excluded.invoice_date,
                seller_company_name = excluded.seller_company_name,
                seller_address = excluded.seller_address,
                seller_gstin = excluded.seller_gstin,
                seller_state = excluded.seller_state,
                seller_state_code = excluded.seller_state_code,
                buyer_name = excluded.buyer_name,
                buyer_address = excluded.buyer_address,
                buyer_gstin = excluded.buyer_gstin,
                buyer_state = excluded.buyer_state,
                buyer_state_code = excluded.buyer_state_code,
                ship_to_name = excluded.ship_to_name,
                ship_to_address = excluded.ship_to_address,
                ship_to_gstin = excluded.ship_to_gstin,
                ship_to_state = excluded.ship_to_state,
                ship_to_state_code = excluded.ship_to_state_code,
                delivery_note = excluded.delivery_note,
                payment_terms = excluded.payment_terms,
                dispatch_through = excluded.dispatch_through,
                destination = excluded.destination,
                vehicle_no = excluded.vehicle_no,
                terms_of_delivery = excluded.terms_of_delivery,
                subtotal = excluded.subtotal,
                cgst = excluded.cgst,
                sgst = excluded.sgst,
                igst = excluded.igst,
                round_off = excluded.round_off,
                grand_total = excluded.grand_total,
                total_quantity = excluded.total_quantity,
                amount_in_words = excluded.amount_in_words,
                bank_details_json = excluded.bank_details_json,
                declaration = excluded.declaration,
                raw_json = excluded.raw_json,
                updated_at = excluded.updated_at
            """,
            {
                "id": invoice.id,
                "invoice_no": invoice.invoice_no,
                "invoice_date": invoice.invoice_date,
                "seller_company_name": invoice.seller.name,
                "seller_address": invoice.seller.address,
                "seller_gstin": invoice.seller.gstin,
                "seller_state": invoice.seller.state,
                "seller_state_code": invoice.seller.state_code,
                "buyer_name": invoice.buyer.name,
                "buyer_address": invoice.buyer.address,
                "buyer_gstin": invoice.buyer.gstin,
                "buyer_state": invoice.buyer.state,
                "buyer_state_code": invoice.buyer.state_code,
                "ship_to_name": invoice.ship_to.name,
                "ship_to_address": invoice.ship_to.address,
                "ship_to_gstin": invoice.ship_to.gstin,
                "ship_to_state": invoice.ship_to.state,
                "ship_to_state_code": invoice.ship_to.state_code,
                "delivery_note": invoice.delivery_note,
                "payment_terms": invoice.payment_terms,
                "dispatch_through": invoice.dispatch_through,
                "destination": invoice.destination,
                "vehicle_no": invoice.vehicle_no,
                "terms_of_delivery": invoice.terms_of_delivery,
                "subtotal": invoice.subtotal,
                "cgst": invoice.cgst,
                "sgst": invoice.sgst,
                "igst": invoice.igst,
                "round_off": invoice.round_off,
                "grand_total": invoice.grand_total,
                "total_quantity": invoice.total_quantity,
                "amount_in_words": invoice.amount_in_words,
                "bank_details_json": json.dumps(payload.get("bank_details", {})),
                "declaration": invoice.declaration,
                "raw_json": json.dumps(payload),
                "created_at": created_at,
                "updated_at": now,
            },
        )
        connection.execute("DELETE FROM invoice_items WHERE invoice_id = ?", (invoice.id,))
        connection.executemany(
            """
            INSERT INTO invoice_items (
                invoice_id, sl_no, description, hsn_sac, quantity, unit, rate, amount
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            [
                (
                    invoice.id,
                    item.sl_no,
                    item.description,
                    item.hsn_sac,
                    item.quantity,
                    item.unit,
                    item.rate,
                    item.amount,
                )
                for item in invoice.items
            ],
        )

    return dump_model(invoice)


def stored_invoice_payloads() -> list[dict]:
    with db_connection() as connection:
        rows = connection.execute(
            "SELECT raw_json FROM invoices ORDER BY updated_at DESC, invoice_date DESC, invoice_no DESC"
        ).fetchall()
    invoices: list[dict] = []
    for row in rows:
        try:
            invoices.append(json.loads(row["raw_json"]))
        except json.JSONDecodeError:
            continue
    return invoices


def clean_text(value: object = "") -> str:
    text = str(value or "")
    replacements = {
        "₹": "Rs. ",
        "–": "-",
        "—": "-",
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',
        "\u00a0": " ",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text.encode("latin-1", "replace").decode("latin-1")


def capacity_kw(quotation: Quotation) -> str:
    value = str(quotation.capacity_kw or "").strip()
    if not value:
        return "0KW"
    return value.upper() if value.upper().endswith("KW") else f"{value}KW"


def rupees(value: str) -> str:
    raw = str(value or "0").strip().replace(",", "").replace("Rs.", "").replace("Rs", "")
    try:
        return f"Rs. {float(raw):,.0f}"
    except ValueError:
        return f"Rs. {value}" if value else "Rs. 0"


def final_price_for_row(value: str) -> str:
    return f"{rupees(value)} /-"


def export_name(quotation: Quotation, extension: str) -> str:
    return f"{quotation.quotation_number}.{extension}"


def invoice_export_name(invoice: Invoice, extension: str) -> str:
    safe_name = "".join(char if char.isalnum() or char in ("-", "_") else "-" for char in invoice.invoice_no)
    return f"{safe_name or 'invoice'}.{extension}"


def quotation_dump(quotation: Quotation) -> dict:
    return dump_model(quotation)


def paragraph_blocks(text: str) -> list[str]:
    return [block.strip() for block in clean_text(text).splitlines() if block.strip()]


def wrap_line(text: str, width: float, font_name: str, font_size: float) -> list[str]:
    words = clean_text(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if pdfmetrics.stringWidth(trial, font_name, font_size) <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_text(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    size: float = 10,
    font: str = "Helvetica",
    fill=colors.black,
) -> None:
    c.setFillColor(fill)
    c.setFont(font, size)
    c.drawString(x, y, clean_text(text))


def draw_centered(
    c: canvas.Canvas,
    text: str,
    y: float,
    size: float = 12,
    font: str = "Helvetica-Bold",
    fill=colors.black,
) -> None:
    c.setFillColor(fill)
    c.setFont(font, size)
    c.drawCentredString(PAGE_WIDTH / 2, y, clean_text(text))


def draw_right(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    size: float = 10,
    font: str = "Helvetica",
    fill=colors.black,
) -> None:
    c.setFillColor(fill)
    c.setFont(font, size)
    c.drawRightString(x, y, clean_text(text))


def draw_wrapped(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    width: float,
    size: float = 10,
    leading: float = 13,
    font: str = "Helvetica",
    fill=colors.black,
) -> float:
    for line in wrap_line(text, width, font, size):
        draw_text(c, line, x, y, size, font, fill)
        y -= leading
    return y


def draw_logo(c: canvas.Canvas, x: float, y: float, width: float, height: float) -> None:
    c.setStrokeColor(LINE)
    c.setFillColor(colors.white)
    c.roundRect(x, y, width, height, 5, stroke=1, fill=1)
    if LOGO_PATH.exists():
        c.drawImage(
            str(LOGO_PATH),
            x + 6,
            y + 6,
            width=width - 12,
            height=height - 12,
            preserveAspectRatio=True,
            mask="auto",
        )
    else:
        c.setFillColor(MUTED)
        c.setFont("Helvetica-Bold", 11)
        c.drawCentredString(x + width / 2, y + height / 2 - 5, "LOGO")


def draw_footer(c: canvas.Canvas) -> None:
    y = 46
    c.setStrokeColor(LINE)
    c.line(MARGIN, y + 17, PAGE_WIDTH - MARGIN, y + 17)
    draw_centered(c, COMPANY_NAME, y + 2, 8, "Helvetica-Bold", DARK)
    draw_centered(c, f"OFFICE ADD. {OFFICE_ADDRESS}", y - 10, 7.2, "Helvetica", MUTED)
    draw_centered(c, f"Email: {EMAIL}   Gst No: {GST_FOOTER}", y - 22, 7.2, "Helvetica", MUTED)


def draw_page_title(c: canvas.Canvas, title: str) -> None:
    draw_logo(c, MARGIN, PAGE_HEIGHT - 94, 94, 58)
    draw_right(c, COMPANY_NAME, PAGE_WIDTH - MARGIN, PAGE_HEIGHT - 58, 14, "Helvetica-Bold", GREEN)
    draw_right(c, title, PAGE_WIDTH - MARGIN, PAGE_HEIGHT - 76, 12, "Helvetica-Bold", DARK)
    c.setStrokeColor(GREEN)
    c.setLineWidth(1.2)
    c.line(MARGIN, PAGE_HEIGHT - 108, PAGE_WIDTH - MARGIN, PAGE_HEIGHT - 108)


def draw_table(
    c: canvas.Canvas,
    x: float,
    y: float,
    col_widths: list[float],
    rows: list[list[str]],
    header_fill=GREEN,
    font_size: float = 8,
    leading: float = 9.5,
    min_height: float = 24,
) -> float:
    for row_index, row in enumerate(rows):
        wrapped_cells = [
            wrap_line(str(row[col_index] if col_index < len(row) else ""), col_widths[col_index] - 8, "Helvetica", font_size)
            for col_index in range(len(col_widths))
        ]
        row_height = max(min_height, max(len(lines) for lines in wrapped_cells) * leading + 10)
        fill = header_fill if row_index == 0 else colors.white
        text_color = colors.white if row_index == 0 else DARK
        font = "Helvetica-Bold" if row_index == 0 else "Helvetica"

        c.setFillColor(fill)
        c.setStrokeColor(LINE)
        c.rect(x, y - row_height, sum(col_widths), row_height, stroke=1, fill=1)

        cursor_x = x
        for col_index, width in enumerate(col_widths):
            if col_index:
                c.line(cursor_x, y, cursor_x, y - row_height)
            text_y = y - 13
            for line in wrapped_cells[col_index]:
                draw_text(c, line, cursor_x + 4, text_y, font_size, font, text_color)
                text_y -= leading
            cursor_x += width
        y -= row_height
    return y


def draw_cover_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_logo(c, MARGIN, PAGE_HEIGHT - 128, 118, 82)

    gst_x = PAGE_WIDTH - MARGIN - 182
    gst_y = PAGE_HEIGHT - 78
    c.setFillColor(SOFT_GOLD)
    c.setStrokeColor(GOLD)
    c.roundRect(gst_x, gst_y, 182, 32, 4, stroke=1, fill=1)
    c.setFillColor(DARK)
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(gst_x + 91, gst_y + 10, f"GST No: {GST_NO}")
    draw_right(c, COMPANY_NAME, PAGE_WIDTH - MARGIN, PAGE_HEIGHT - 112, 17, "Helvetica-Bold", GREEN)

    draw_centered(c, "QUOTATION", PAGE_HEIGHT - 225, 28, "Helvetica-Bold", DARK)
    draw_centered(c, "Proposal for installation, Engineering and Procurement of", PAGE_HEIGHT - 270, 13, "Helvetica-Bold", DARK)
    draw_centered(c, "ON GRID Connected Roof-Top Solar PV System.", PAGE_HEIGHT - 290, 13, "Helvetica-Bold", DARK)

    box_w = PAGE_WIDTH - (MARGIN * 2) - 80
    x = MARGIN + 40
    y = PAGE_HEIGHT - 410
    c.setStrokeColor(LINE)
    c.setFillColor(colors.white)
    c.roundRect(x, y - 142, box_w, 142, 8, stroke=1, fill=1)
    details = [
        ("NAME -", quotation.customer_name),
        ("LOCATION -", quotation.location),
        ("CAPACITY -", capacity_kw(quotation)),
        ("DATE -", quotation.quote_date),
    ]
    row_y = y - 30
    for label, value in details:
        draw_text(c, label, x + 28, row_y, 12, "Helvetica-Bold", DARK)
        draw_text(c, value, x + 130, row_y, 12, "Helvetica", DARK)
        row_y -= 28

    draw_footer(c)


def draw_introduction_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_page_title(c, "INTRODUCTION")
    y = PAGE_HEIGHT - 142
    for block in paragraph_blocks(quotation.introduction):
        if block.lower().startswith("why to go solar"):
            y -= 8
            draw_text(c, block, MARGIN, y, 13, "Helvetica-Bold", GREEN)
            y -= 22
            continue
        y = draw_wrapped(c, block, MARGIN, y, PAGE_WIDTH - (MARGIN * 2), 9.6, 13, "Helvetica")
        y -= 7
    draw_footer(c)


def draw_specification_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_page_title(c, "SYSTEM SPECIFICATION")
    rows = [["ITEM", "BRAND/PARAMETERS", "DETAIL", "SPECIFICATION", "QUANTITY"]]
    rows.extend(
        [
            [spec.item, spec.brand_parameters, spec.detail, spec.specification, spec.quantity]
            for spec in quotation.specifications
        ]
    )
    draw_table(c, MARGIN, PAGE_HEIGHT - 142, [92, 126, 116, 108, 80], rows, font_size=7.2, leading=8.5, min_height=28)
    draw_footer(c)


def draw_price_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_page_title(c, "PRICE QUOTION")
    y = PAGE_HEIGHT - 150
    c.setFillColor(SOFT_GREEN)
    c.setStrokeColor(GREEN)
    c.roundRect(MARGIN, y - 40, PAGE_WIDTH - (MARGIN * 2), 40, 5, stroke=1, fill=1)
    draw_centered(c, quotation.price_header or f"{capacity_kw(quotation)} SANCTION AND {capacity_kw(quotation)} PV MODULE", y - 16, 12, "Helvetica-Bold", DARK)
    y -= 64

    rows = [["SR NO", "SYSTEM DETAIL", "RATE"]]
    for row in quotation.price_rows:
        rate = row.rate
        if row.system_detail.strip().upper() == "TOTAL":
            rate = final_price_for_row(quotation.final_price)
        rows.append([row.sr_no, row.system_detail, rate])
    y = draw_table(c, MARGIN + 34, y, [80, 260, 126], rows, font_size=9, leading=11, min_height=30)
    y -= 28
    draw_wrapped(c, f"* {quotation.price_note}", MARGIN + 20, y, PAGE_WIDTH - (MARGIN * 2) - 40, 10, 13, "Helvetica-Bold", DARK)
    y -= 72
    c.setFillColor(SOFT_GOLD)
    c.setStrokeColor(GOLD)
    c.roundRect(MARGIN, y - 50, PAGE_WIDTH - (MARGIN * 2), 50, 5, stroke=1, fill=1)
    draw_centered(c, quotation.payment_note, y - 21, 10, "Helvetica-Bold", DARK)
    draw_footer(c)


def draw_terms_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_page_title(c, "TERMS & CONDITIONS")
    columns = [
        (MARGIN, PAGE_HEIGHT - 136, (PAGE_WIDTH - (MARGIN * 2) - 18) / 2),
        (MARGIN + (PAGE_WIDTH - (MARGIN * 2) + 18) / 2, PAGE_HEIGHT - 136, (PAGE_WIDTH - (MARGIN * 2) - 18) / 2),
    ]
    column_index = 0
    x, y, width = columns[column_index]
    bottom = 78

    for block in paragraph_blocks(quotation.terms_conditions):
        font = "Helvetica-Bold" if block.upper().startswith("TERMS") else "Helvetica"
        size = 6.8 if font == "Helvetica" else 7.6
        leading = 8.2
        lines = wrap_line(block, width, font, size)
        needed = len(lines) * leading + 5
        if y - needed < bottom and column_index == 0:
            column_index = 1
            x, y, width = columns[column_index]
        for line in lines:
            if y < bottom:
                break
            draw_text(c, line, x, y, size, font, DARK)
            y -= leading
        y -= 5
    draw_footer(c)


def draw_bank_page(c: canvas.Canvas, quotation: Quotation) -> None:
    draw_page_title(c, "BANK DETAILS + SIGNATURE")
    y = PAGE_HEIGHT - 152
    c.setFillColor(colors.white)
    c.setStrokeColor(LINE)
    c.roundRect(MARGIN, y - 92, PAGE_WIDTH - (MARGIN * 2), 92, 6, stroke=1, fill=1)
    draw_text(c, "REMARK IF ANY", MARGIN + 18, y - 24, 11, "Helvetica-Bold", DARK)
    draw_wrapped(c, quotation.remark or "-", MARGIN + 18, y - 48, PAGE_WIDTH - (MARGIN * 2) - 36, 10, 13, "Helvetica")

    y -= 132
    bank_rows = [
        ["MAIL ID", EMAIL],
        ["CONTACT NO", CONTACT_NO],
        ["ACCOUNT HOLDER", ACCOUNT_HOLDER],
        ["BANK NAME", BANK_NAME],
        ["CC ACCOUNT NO", CC_ACCOUNT_NO],
        ["IFSC CODE", IFSC_CODE],
        ["GST", GST_NO],
        ["BRANCH", BRANCH],
    ]
    draw_table(c, MARGIN + 30, y, [170, 310], [["PARTICULAR", "DETAIL"], *bank_rows], font_size=9, leading=11, min_height=28)

    sig_y = 154
    c.setStrokeColor(DARK)
    c.line(PAGE_WIDTH - 255, sig_y, PAGE_WIDTH - MARGIN, sig_y)
    draw_right(c, "SIGNATURE & SEAL OF DIRECTOR", PAGE_WIDTH - MARGIN, sig_y - 18, 10, "Helvetica-Bold", DARK)
    draw_footer(c)


def create_pdf(quotation: Quotation) -> bytes:
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=A4)

    draw_cover_page(c, quotation)
    c.showPage()
    draw_introduction_page(c, quotation)
    c.showPage()
    draw_specification_page(c, quotation)
    c.showPage()
    draw_price_page(c, quotation)
    c.showPage()
    draw_terms_page(c, quotation)
    c.showPage()
    draw_bank_page(c, quotation)
    c.save()

    return output.getvalue()


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)


def add_doc_title(document: Document, title: str) -> None:
    if LOGO_PATH.exists():
        document.add_picture(str(LOGO_PATH), width=Inches(1.35))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = heading.add_run(COMPANY_NAME)
    run.bold = True
    run.font.size = Pt(16)
    sub = document.add_paragraph(title)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(12)


def add_doc_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for index, header in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = clean_text(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = clean_text(value)


def add_bold_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(clean_text(text))
    run.bold = True


def add_page_break(document: Document) -> None:
    document.add_page_break()


def create_word(quotation: Quotation) -> bytes:
    document = Document()
    set_doc_defaults(document)

    add_doc_title(document, "QUOTATION")
    document.add_paragraph(
        "Proposal for installation, Engineering and Procurement of ON GRID Connected Roof-Top Solar PV System."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_doc_table(
        document,
        [
            ["Field", "Details"],
            ["Quotation Number", quotation.quotation_number],
            ["NAME", quotation.customer_name],
            ["LOCATION", quotation.location],
            ["CAPACITY", capacity_kw(quotation)],
            ["DATE", quotation.quote_date],
            ["GST No", GST_NO],
        ],
    )

    add_page_break(document)
    add_doc_title(document, "INTRODUCTION")
    for block in paragraph_blocks(quotation.introduction):
        paragraph = document.add_paragraph(clean_text(block))
        if block.lower().startswith("why to go solar"):
            paragraph.runs[0].bold = True

    add_page_break(document)
    add_doc_title(document, "SYSTEM SPECIFICATION")
    add_doc_table(
        document,
        [
            ["ITEM", "BRAND/PARAMETERS", "DETAIL", "SPECIFICATION", "QUANTITY"],
            *[
                [spec.item, spec.brand_parameters, spec.detail, spec.specification, spec.quantity]
                for spec in quotation.specifications
            ],
        ],
    )

    add_page_break(document)
    add_doc_title(document, "PRICE QUOTION")
    add_bold_paragraph(document, quotation.price_header)
    add_doc_table(
        document,
        [
            ["SR NO", "SYSTEM DETAIL", "RATE"],
            *[
                [
                    row.sr_no,
                    row.system_detail,
                    final_price_for_row(quotation.final_price)
                    if row.system_detail.strip().upper() == "TOTAL"
                    else row.rate,
                ]
                for row in quotation.price_rows
            ],
        ],
    )
    document.add_paragraph(f"* {clean_text(quotation.price_note)}")
    add_bold_paragraph(document, quotation.payment_note)

    add_page_break(document)
    add_doc_title(document, "TERMS & CONDITIONS")
    for block in paragraph_blocks(quotation.terms_conditions):
        document.add_paragraph(clean_text(block))

    add_page_break(document)
    add_doc_title(document, "BANK DETAILS + SIGNATURE")
    add_bold_paragraph(document, "REMARK IF ANY")
    document.add_paragraph(clean_text(quotation.remark or "-"))
    add_doc_table(
        document,
        [
            ["PARTICULAR", "DETAIL"],
            ["MAIL ID", EMAIL],
            ["CONTACT NO", CONTACT_NO],
            ["ACCOUNT HOLDER", ACCOUNT_HOLDER],
            ["BANK NAME", BANK_NAME],
            ["CC ACCOUNT NO", CC_ACCOUNT_NO],
            ["IFSC CODE", IFSC_CODE],
            ["GST", GST_NO],
            ["BRANCH", BRANCH],
        ],
    )
    document.add_paragraph("\n\nSIGNATURE & SEAL OF DIRECTOR").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def write_sheet(ws, rows: list[list[object]]) -> None:
    header_fill = PatternFill("solid", fgColor="176B3A")
    for row in rows:
        ws.append([clean_text(value) for value in row])
    if ws.max_row:
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for column in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column)
        ws.column_dimensions[column[0].column_letter].width = min(max_length + 4, 48)


def create_excel(quotation: Quotation) -> bytes:
    workbook = Workbook()
    cover = workbook.active
    cover.title = "01 Cover Page"
    write_sheet(
        cover,
        [
            ["Field", "Details"],
            ["Company", COMPANY_NAME],
            ["Quotation Number", quotation.quotation_number],
            ["Customer Name", quotation.customer_name],
            ["Location", quotation.location],
            ["Capacity", capacity_kw(quotation)],
            ["Date", quotation.quote_date],
            ["GST No", GST_NO],
            ["Office Address", OFFICE_ADDRESS],
            ["Email", EMAIL],
        ],
    )
    write_sheet(workbook.create_sheet("02 Introduction"), [["Introduction"], *[[block] for block in paragraph_blocks(quotation.introduction)]])
    write_sheet(
        workbook.create_sheet("03 Specification"),
        [
            ["ITEM", "BRAND/PARAMETERS", "DETAIL", "SPECIFICATION", "QUANTITY"],
            *[
                [spec.item, spec.brand_parameters, spec.detail, spec.specification, spec.quantity]
                for spec in quotation.specifications
            ],
        ],
    )
    write_sheet(
        workbook.create_sheet("04 Price"),
        [
            ["PRICE QUOTION", quotation.price_header],
            ["SR NO", "SYSTEM DETAIL", "RATE"],
            *[
                [
                    row.sr_no,
                    row.system_detail,
                    final_price_for_row(quotation.final_price)
                    if row.system_detail.strip().upper() == "TOTAL"
                    else row.rate,
                ]
                for row in quotation.price_rows
            ],
            ["NOTE", quotation.price_note],
            ["PAYMENT NOTE", quotation.payment_note],
        ],
    )
    write_sheet(workbook.create_sheet("05 Terms"), [["Terms & Conditions"], *[[block] for block in paragraph_blocks(quotation.terms_conditions)]])
    write_sheet(
        workbook.create_sheet("06 Bank Details"),
        [
            ["PARTICULAR", "DETAIL"],
            ["REMARK IF ANY", quotation.remark or "-"],
            ["SIGNATURE & SEAL", "SIGNATURE & SEAL OF DIRECTOR"],
            ["MAIL ID", EMAIL],
            ["CONTACT NO", CONTACT_NO],
            ["ACCOUNT HOLDER", ACCOUNT_HOLDER],
            ["BANK NAME", BANK_NAME],
            ["CC ACCOUNT NO", CC_ACCOUNT_NO],
            ["IFSC CODE", IFSC_CODE],
            ["GST", GST_NO],
            ["BRANCH", BRANCH],
        ],
    )

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def money(value: float | int | str) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


def tax_rate(value: float | int | str, fallback: float) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return float(fallback)
    return number if number >= 0 else float(fallback)


def tax_label(label: str, value: float | int | str, fallback: float) -> str:
    rate = tax_rate(value, fallback)
    if rate.is_integer():
        rate_text = str(int(rate))
    else:
        rate_text = f"{rate:.2f}".rstrip("0").rstrip(".")
    return f"{label} {rate_text}%"


def number_text(value: float | int | str) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "0"
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}".rstrip("0").rstrip(".")


def invoice_tax_summary(invoice: Invoice) -> list[InvoiceTaxSummaryRow]:
    if invoice.tax_summary:
        return invoice.tax_summary

    cgst_rate = tax_rate(invoice.cgst_rate, 9) / 100
    sgst_rate = tax_rate(invoice.sgst_rate, 9) / 100
    igst_rate = tax_rate(invoice.igst_rate, 18) / 100
    grouped: dict[str, float] = {}
    for item in invoice.items:
        key = item.hsn_sac.strip() or "-"
        grouped[key] = grouped.get(key, 0) + float(item.amount or 0)

    rows: list[InvoiceTaxSummaryRow] = []
    for hsn_sac, taxable in grouped.items():
        if invoice.use_igst:
            igst = round(taxable * igst_rate, 2)
            rows.append(
                InvoiceTaxSummaryRow(
                    hsn_sac=hsn_sac,
                    taxable_value=taxable,
                    igst=igst,
                    total_tax=igst,
                )
            )
        else:
            cgst = round(taxable * cgst_rate, 2)
            sgst = round(taxable * sgst_rate, 2)
            rows.append(
                InvoiceTaxSummaryRow(
                    hsn_sac=hsn_sac,
                    taxable_value=taxable,
                    cgst=cgst,
                    sgst=sgst,
                    total_tax=cgst + sgst,
                )
            )
    return rows


def draw_invoice_cell_row(
    c: canvas.Canvas,
    x: float,
    y: float,
    col_widths: list[float],
    cells: list[str],
    row_height: float,
    fill=colors.white,
    text_color=DARK,
    font: str = "Helvetica",
    font_size: float = 7.8,
    leading: float = 9.2,
) -> float:
    c.setStrokeColor(LINE)
    c.setFillColor(fill)
    c.rect(x, y - row_height, sum(col_widths), row_height, stroke=1, fill=1)

    cursor_x = x
    for index, width in enumerate(col_widths):
        if index:
            c.line(cursor_x, y, cursor_x, y - row_height)
        text_y = y - 12
        for line in wrap_line(str(cells[index] if index < len(cells) else ""), width - 8, font, font_size):
            draw_text(c, line, cursor_x + 4, text_y, font_size, font, text_color)
            text_y -= leading
        cursor_x += width
    return y - row_height


def invoice_item_row_height(cells: list[str], col_widths: list[float], font_size: float = 7.8, leading: float = 9.2) -> float:
    line_counts = [
        len(wrap_line(str(cells[index] if index < len(cells) else ""), width - 8, "Helvetica", font_size))
        for index, width in enumerate(col_widths)
    ]
    return max(24, max(line_counts or [1]) * leading + 10)


def draw_invoice_header(c: canvas.Canvas, invoice: Invoice, page_label: str = "") -> float:
    seller = invoice.seller
    c.setTitle(clean_text(f"{invoice.invoice_no} - Tax Invoice"))
    c.setStrokeColor(DARK)
    c.setLineWidth(0.8)
    c.rect(MARGIN, 28, PAGE_WIDTH - (MARGIN * 2), PAGE_HEIGHT - 56, stroke=1, fill=0)

    draw_centered(c, "TAX INVOICE / SOLAR INVOICE", PAGE_HEIGHT - 42, 13, "Helvetica-Bold", DARK)
    if page_label:
        draw_right(c, page_label, PAGE_WIDTH - MARGIN - 8, PAGE_HEIGHT - 42, 8.5, "Helvetica-Bold", MUTED)

    header_top = PAGE_HEIGHT - 56
    header_bottom = PAGE_HEIGHT - 144
    c.setStrokeColor(LINE)
    c.line(MARGIN, header_bottom, PAGE_WIDTH - MARGIN, header_bottom)
    c.line(PAGE_WIDTH - 222, header_top, PAGE_WIDTH - 222, header_bottom)

    draw_logo(c, MARGIN + 8, header_bottom + 14, 66, 52)
    seller_x = MARGIN + 84
    draw_text(c, seller.name or COMPANY_NAME, seller_x, header_top - 20, 12, "Helvetica-Bold", GREEN)
    y = header_top - 36
    for line in wrap_line(seller.address or OFFICE_ADDRESS, PAGE_WIDTH - MARGIN - 236 - seller_x, "Helvetica", 7.8):
        draw_text(c, line, seller_x, y, 7.8, "Helvetica", DARK)
        y -= 9.2
    draw_text(c, f"GSTIN/UIN: {seller.gstin or GST_NO}", seller_x, y - 2, 8, "Helvetica-Bold", DARK)
    draw_text(
        c,
        f"State: {seller.state or 'Maharashtra'}   Code: {seller.state_code or '27'}",
        seller_x,
        y - 14,
        8,
        "Helvetica",
        DARK,
    )

    invoice_x = PAGE_WIDTH - 214
    details = [
        ("Invoice No.", invoice.invoice_no),
        ("Date", invoice.invoice_date),
        ("Delivery Note", invoice.delivery_note),
        ("Payment Terms", invoice.payment_terms),
    ]
    detail_y = header_top - 18
    for label, value in details:
        draw_text(c, label, invoice_x, detail_y, 7.5, "Helvetica-Bold", MUTED)
        draw_text(c, value or "-", invoice_x + 78, detail_y, 7.5, "Helvetica", DARK)
        detail_y -= 16

    return header_bottom - 14


def draw_invoice_parties(c: canvas.Canvas, invoice: Invoice, y: float) -> float:
    box_height = 108
    col_width = (PAGE_WIDTH - (MARGIN * 2)) / 2
    c.setStrokeColor(LINE)
    c.rect(MARGIN, y - box_height, col_width, box_height, stroke=1, fill=0)
    c.rect(MARGIN + col_width, y - box_height, col_width, box_height, stroke=1, fill=0)

    def draw_party(title: str, party: InvoiceParty, x: float) -> None:
        draw_text(c, title, x + 8, y - 14, 8.3, "Helvetica-Bold", MUTED)
        draw_text(c, party.name or "-", x + 8, y - 30, 9.2, "Helvetica-Bold", DARK)
        text_y = y - 44
        for line in wrap_line(party.address or "-", col_width - 18, "Helvetica", 7.7):
            draw_text(c, line, x + 8, text_y, 7.7, "Helvetica", DARK)
            text_y -= 9
        draw_text(c, f"GSTIN/UIN: {party.gstin or '-'}", x + 8, y - 86, 7.8, "Helvetica-Bold", DARK)
        draw_text(c, f"State: {party.state or '-'}   Code: {party.state_code or '-'}", x + 8, y - 99, 7.8, "Helvetica", DARK)

    draw_party("Buyer (Bill To)", invoice.buyer, MARGIN)
    draw_party("Ship To", invoice.ship_to, MARGIN + col_width)
    return y - box_height - 12


def draw_dispatch_details(c: canvas.Canvas, invoice: Invoice, y: float) -> float:
    rows = [
        ["Dispatch Through", invoice.dispatch_through or "-", "Destination", invoice.destination or "-"],
        ["Vehicle No.", invoice.vehicle_no or "-", "Terms of Delivery", invoice.terms_of_delivery or "-"],
    ]
    return draw_table(c, MARGIN, y, [112, 150, 112, 149], rows, header_fill=GREEN, font_size=7.8, leading=9.2, min_height=24) - 12


def draw_invoice_items(c: canvas.Canvas, invoice: Invoice, y: float) -> float:
    col_widths = [34, 196, 64, 42, 44, 66, 77]
    header = ["Sl No", "Description", "HSN/SAC", "Qty", "Unit", "Rate", "Amount"]
    y = draw_invoice_cell_row(c, MARGIN, y, col_widths, header, 24, fill=GREEN, text_color=colors.white, font="Helvetica-Bold")

    items = invoice.items or [InvoiceItem()]
    for index, item in enumerate(items, start=1):
        cells = [
            item.sl_no or str(index),
            item.description,
            item.hsn_sac,
            number_text(item.quantity),
            item.unit,
            money(item.rate),
            money(item.amount),
        ]
        row_height = invoice_item_row_height(cells, col_widths)
        if y - row_height < 244:
            draw_text(c, "Continued on next page", MARGIN + 8, 46, 8, "Helvetica-Bold", MUTED)
            c.showPage()
            y = draw_invoice_header(c, invoice, "Continued")
            y = draw_invoice_cell_row(c, MARGIN, y, col_widths, header, 24, fill=GREEN, text_color=colors.white, font="Helvetica-Bold")
        y = draw_invoice_cell_row(c, MARGIN, y, col_widths, cells, row_height)
    return y - 12


def draw_invoice_totals(c: canvas.Canvas, invoice: Invoice, y: float) -> float:
    if y < 268:
        c.showPage()
        y = draw_invoice_header(c, invoice, "Totals")

    cgst_label = tax_label("CGST", invoice.cgst_rate, 9)
    sgst_label = tax_label("SGST", invoice.sgst_rate, 9)
    igst_label = tax_label("IGST", invoice.igst_rate, 18)
    amount_words = invoice.amount_in_words or f"Rupees {money(invoice.grand_total)} Only"
    c.setFillColor(SOFT_GOLD)
    c.setStrokeColor(GOLD)
    c.rect(MARGIN, y - 42, PAGE_WIDTH - (MARGIN * 2), 42, stroke=1, fill=1)
    draw_text(c, "Amount in Words", MARGIN + 8, y - 14, 7.8, "Helvetica-Bold", DARK)
    draw_wrapped(c, amount_words, MARGIN + 8, y - 28, PAGE_WIDTH - (MARGIN * 2) - 16, 8.2, 9.4, "Helvetica-Bold", DARK)
    y -= 54

    tax_rows = [["HSN/SAC", "Taxable", cgst_label, sgst_label, igst_label, "Tax Total"]]
    for row in invoice_tax_summary(invoice):
        tax_rows.append(
            [
                row.hsn_sac or "-",
                money(row.taxable_value),
                money(row.cgst),
                money(row.sgst),
                money(row.igst),
                money(row.total_tax),
            ]
        )
    if len(tax_rows) == 1:
        tax_rows.append(["-", "0.00", "0.00", "0.00", "0.00", "0.00"])

    tax_bottom = draw_table(c, MARGIN, y, [72, 86, 78, 78, 78, 96], tax_rows, font_size=7.3, leading=8.6, min_height=22)
    y = tax_bottom - 10

    totals = [
        ["Total Qty", number_text(invoice.total_quantity)],
        ["Subtotal", money(invoice.subtotal)],
        [cgst_label, money(invoice.cgst)],
        [sgst_label, money(invoice.sgst)],
        [igst_label, money(invoice.igst)],
        ["Round Off", money(invoice.round_off)],
        ["Grand Total", money(invoice.grand_total)],
    ]
    totals_bottom = draw_table(
        c,
        PAGE_WIDTH - MARGIN - 204,
        y,
        [104, 100],
        totals,
        header_fill=GREEN,
        font_size=7.8,
        leading=9.2,
        min_height=22,
    )
    return totals_bottom - 12


def draw_invoice_footer(c: canvas.Canvas, invoice: Invoice, y: float) -> None:
    if y < 158:
        c.showPage()
        y = draw_invoice_header(c, invoice, "Bank Details")

    bank = invoice.bank_details
    bank_rows = [
        ["Account Holder", bank.account_holder or ACCOUNT_HOLDER],
        ["Bank Name", bank.bank_name or BANK_NAME],
        ["Account No", bank.account_no or CC_ACCOUNT_NO],
        ["IFSC", bank.ifsc or IFSC_CODE],
        ["Branch", bank.branch or BRANCH],
    ]
    c.setStrokeColor(LINE)
    c.rect(MARGIN, y - 116, 286, 116, stroke=1, fill=0)
    draw_text(c, "Bank Details", MARGIN + 8, y - 14, 8.5, "Helvetica-Bold", DARK)
    detail_y = y - 30
    for label, value in bank_rows:
        draw_text(c, label, MARGIN + 8, detail_y, 7.4, "Helvetica-Bold", MUTED)
        draw_text(c, value, MARGIN + 98, detail_y, 7.4, "Helvetica", DARK)
        detail_y -= 16

    declaration_x = MARGIN + 300
    declaration_w = PAGE_WIDTH - MARGIN - declaration_x
    c.rect(declaration_x, y - 116, declaration_w, 116, stroke=1, fill=0)
    draw_text(c, "Declaration", declaration_x + 8, y - 14, 8.5, "Helvetica-Bold", DARK)
    draw_wrapped(
        c,
        invoice.declaration or "We declare that this invoice shows the actual price of the goods and services described and that all particulars are true and correct.",
        declaration_x + 8,
        y - 30,
        declaration_w - 16,
        7.4,
        8.8,
        "Helvetica",
    )
    draw_right(c, f"For {invoice.seller.name or COMPANY_NAME}", PAGE_WIDTH - MARGIN - 8, y - 74, 8.2, "Helvetica-Bold", DARK)
    draw_right(c, "Authorised Signatory", PAGE_WIDTH - MARGIN - 8, y - 102, 7.8, "Helvetica-Bold", DARK)


def create_invoice_pdf(invoice: Invoice) -> bytes:
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    y = draw_invoice_header(c, invoice)
    y = draw_invoice_parties(c, invoice, y)
    y = draw_dispatch_details(c, invoice, y)
    y = draw_invoice_items(c, invoice, y)
    y = draw_invoice_totals(c, invoice, y)
    draw_invoice_footer(c, invoice, y)
    c.save()
    return output.getvalue()


init_db()


@app.post("/save-quotation")
def save_quotation(quotation: Quotation) -> dict:
    global saved_quotations
    saved_quotations = [item for item in saved_quotations if item.id != quotation.id]
    saved_quotations.insert(0, quotation)
    return {"success": True, "message": "Quotation accepted. Browser localStorage remains the storage for V1."}


@app.get("/history")
def history() -> dict:
    return {"quotations": [quotation_dump(item) for item in saved_quotations]}


@app.post("/save-invoice")
def save_invoice(invoice: Invoice) -> dict:
    saved_invoice = save_invoice_to_db(invoice)
    return {"success": True, "message": "Invoice saved to database.", "invoice": saved_invoice}


@app.get("/invoices")
def invoices() -> dict:
    return {"invoices": stored_invoice_payloads()}


@app.post("/generate-pdf")
def generate_pdf(quotation: Quotation) -> Response:
    content = create_pdf(quotation)
    return Response(
        content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{export_name(quotation, "pdf")}"'},
    )


@app.post("/generate-word")
def generate_word(quotation: Quotation) -> Response:
    content = create_word(quotation)
    return Response(
        content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{export_name(quotation, "docx")}"'},
    )


@app.post("/generate-excel")
def generate_excel(quotation: Quotation) -> Response:
    content = create_excel(quotation)
    return Response(
        content,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{export_name(quotation, "xlsx")}"'},
    )


@app.post("/generate-invoice-pdf")
def generate_invoice_pdf(invoice: Invoice) -> Response:
    content = create_invoice_pdf(invoice)
    return Response(
        content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{invoice_export_name(invoice, "pdf")}"'},
    )


app.mount("/", StaticFiles(directory=str(BASE_DIR), html=True), name="static")
